#!/usr/bin/env python3
"""
Otto Doc-Writer v2：专业排版引擎 — AI 写 Markdown → 设计级 .docx

用法：
  python create_docx.py <input.md> <output.docx> [-p official|report|letter|meeting|proposal]

设计升级（v2）：
  - 封面页：色块 + 大标题 + 元信息
  - Word 原生标题样式（Heading 1–3），可在导航窗格中跳转
  - 真实的 TOC 目录域
  - 专业表格：细线 + 深色表头 + 交替行色
  - 高亮块：左侧 accent bar + 浅底色
  - 页眉装饰线 + 页脚页码
  - 安全字体降级（无指定字体时自动回退系统字体）
"""
from __future__ import annotations

import json, os, re, sys
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

if sys.platform == 'win32':
    try: sys.stdout.reconfigure(encoding='utf-8')
    except Exception: pass

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Emu, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：需要 python-docx。pip install python-docx"); sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════
# 预设模板
# ═══════════════════════════════════════════════════════════════════════

PRESETS = {
    "report": {
        "name": "报告",
        "cover": True,
        "toc": True,
        "page": {"width": Cm(21), "height": Cm(29.7),
                 "margin_top": Cm(2.54), "margin_bottom": Cm(2.54),
                 "margin_left": Cm(3.18), "margin_right": Cm(3.18)},
        "fonts": {"title": "Microsoft YaHei", "heading": "Microsoft YaHei",
                  "body": "SimSun", "body_size": Pt(12),
                  "title_size": Pt(26), "heading_size": Pt(16)},
        "colors": {"primary": "1B3A5C", "accent": "2E75B6", "body": "333333",
                   "cover_bg": "1B3A5C", "cover_text": "FFFFFF",
                   "table_header_bg": "1B3A5C", "table_header_text": "FFFFFF",
                   "table_stripe": "F0F4F8", "callout_bg": "EDF2F9",
                   "callout_bar": "2E75B6", "hr_color": "CCCCCC"},
        "header_line": True,
    },
    "official": {
        "name": "公文",
        "cover": False,
        "toc": False,
        "page": {"width": Cm(21), "height": Cm(29.7),
                 "margin_top": Cm(3.7), "margin_bottom": Cm(3.5),
                 "margin_left": Cm(2.8), "margin_right": Cm(2.6)},
        "fonts": {"title": "FZXiaoBiaoSong-B05S", "heading": "SimHei",
                  "body": "FangSong", "body_size": Pt(16),
                  "title_size": Pt(22), "heading_size": Pt(16)},
        "colors": {"primary": "C00000", "accent": "C00000", "body": "000000",
                   "cover_bg": "", "cover_text": "",
                   "table_header_bg": "C00000", "table_header_text": "FFFFFF",
                   "table_stripe": "FFF5F5", "callout_bg": "FFF8F8",
                   "callout_bar": "C00000", "hr_color": "C00000"},
        "title_align": "center",
        "title_line": True,
        "first_indent": True,
        "line_spacing": 28,  # 固定行距 28 磅
        "header_line": False,
    },
    "proposal": {
        "name": "方案",
        "cover": True,
        "toc": True,
        "page": {"width": Cm(21), "height": Cm(29.7),
                 "margin_top": Cm(2.54), "margin_bottom": Cm(2.54),
                 "margin_left": Cm(2.54), "margin_right": Cm(2.54)},
        "fonts": {"title": "Microsoft YaHei", "heading": "Microsoft YaHei",
                  "body": "Microsoft YaHei", "body_size": Pt(11),
                  "title_size": Pt(28), "heading_size": Pt(15)},
        "colors": {"primary": "2E75B6", "accent": "FF6B35", "body": "333333",
                   "cover_bg": "2E75B6", "cover_text": "FFFFFF",
                   "table_header_bg": "2E75B6", "table_header_text": "FFFFFF",
                   "table_stripe": "F0F7FF", "callout_bg": "FFF8F5",
                   "callout_bar": "FF6B35", "hr_color": "E0E0E0"},
        "header_line": True,
    },
    "letter": {
        "name": "信函",
        "cover": False,
        "toc": False,
        "page": {"width": Cm(21), "height": Cm(29.7),
                 "margin_top": Cm(3), "margin_bottom": Cm(2.5),
                 "margin_left": Cm(3), "margin_right": Cm(3)},
        "fonts": {"title": "SimSun", "heading": "SimHei",
                  "body": "SimSun", "body_size": Pt(12),
                  "title_size": Pt(16), "heading_size": Pt(14)},
        "colors": {"primary": "000000", "accent": "333333", "body": "000000",
                   "cover_bg": "", "cover_text": "",
                   "table_header_bg": "333333", "table_header_text": "FFFFFF",
                   "table_stripe": "F5F5F5", "callout_bg": "F9F9F9",
                   "callout_bar": "999999", "hr_color": "CCCCCC"},
        "title_align": "left",
        "header_line": False,
    },
    "meeting": {
        "name": "会议纪要",
        "cover": False,
        "toc": False,
        "page": {"width": Cm(21), "height": Cm(29.7),
                 "margin_top": Cm(2.54), "margin_bottom": Cm(2.54),
                 "margin_left": Cm(2.54), "margin_right": Cm(2.54)},
        "fonts": {"title": "SimHei", "heading": "SimHei",
                  "body": "FangSong", "body_size": Pt(12),
                  "title_size": Pt(20), "heading_size": Pt(14)},
        "colors": {"primary": "000000", "accent": "555555", "body": "333333",
                   "cover_bg": "", "cover_text": "",
                   "table_header_bg": "4A4A4A", "table_header_text": "FFFFFF",
                   "table_stripe": "F5F5F5", "callout_bg": "FFFCF0",
                   "callout_bar": "D4A843", "hr_color": "CCCCCC"},
        "title_align": "center",
        "header_line": False,
    },
}

# 字体降级映射：如果系统没有首选字体，尝试这些
FONT_FALLBACK = {
    "Microsoft YaHei": ["微软雅黑", "SimHei", "Arial"],
    "SimHei": ["黑体", "Microsoft YaHei", "Arial"],
    "FZXiaoBiaoSong-B05S": ["方正小标宋简体", "SimSun", "宋体", "SimHei"],
    "SimSun": ["宋体", "Microsoft YaHei", "Arial"],
    "FangSong": ["仿宋", "SimSun", "宋体", "Arial"],
}


# ═══════════════════════════════════════════════════════════════════════
# Markdown 解析（与 v1 兼容）
# ═══════════════════════════════════════════════════════════════════════

def parse_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    meta: dict[str, Any] = {}
    body = text.strip()
    if body.startswith("---"):
        parts = body.split("---", 2)
        if len(parts) >= 3:
            for line in parts[1].strip().split("\n"):
                line = line.strip()
                if ":" in line:
                    key, _, val = line.partition(":")
                    meta[key.strip()] = val.strip().strip('"').strip("'")
            body = parts[2].strip()
    return meta, body


def parse_markdown_body(text: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    lines = text.split("\n"); i = 0
    table_buf: list[str] = []; in_table = False

    def flush_table():
        nonlocal table_buf, in_table
        if table_buf:
            blocks.append({"type": "table", "raw": table_buf})
            table_buf = []; in_table = False

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line); in_table = True; i += 1; continue
        elif in_table:
            flush_table(); continue
        if not line.strip(): i += 1; continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            flush_table()
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1; continue

        m = re.match(r"^[-*+]\s+(.+)$", line)
        if m:
            flush_table()
            items = []
            while i < len(lines) and re.match(r"^[-*+]\s+(.+)$", lines[i]):
                items.append(re.match(r"^[-*+]\s+(.+)$", lines[i]).group(1).strip()); i += 1
            blocks.append({"type": "bullet_list", "items": items}); continue

        m = re.match(r"^\d+[.)]\s+(.+)$", line)
        if m:
            flush_table()
            items = []
            while i < len(lines) and re.match(r"^\d+[.)]\s+(.+)$", lines[i]):
                items.append(re.match(r"^\d+[.)]\s+(.+)$", lines[i]).group(1).strip()); i += 1
            blocks.append({"type": "ordered_list", "items": items}); continue

        if line.startswith("> "):
            flush_table()
            qlines = []
            while i < len(lines) and lines[i].startswith("> "):
                qlines.append(lines[i][2:].strip()); i += 1
            blocks.append({"type": "quote", "text": "\n".join(qlines)}); continue

        if line.strip() in ("---", "***", "___"):
            flush_table()
            blocks.append({"type": "hr"}); i += 1; continue

        flush_table()
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and \
              not re.match(r"^[-*+]\s+", lines[i]) and not re.match(r"^\d+[.)]\s+", lines[i]) and \
              not lines[i].startswith("|") and not lines[i].startswith("> ") and \
              lines[i].strip() not in ("---","***","___"):
            para_lines.append(lines[i]); i += 1
        blocks.append({"type": "paragraph", "text": "\n".join(para_lines)})

    flush_table()
    return blocks


def parse_table(raw: list[str]) -> tuple[list[str], list[list[str]]]:
    if len(raw) < 2: return [], []
    headers = [c.strip() for c in raw[0].strip("|").split("|")]
    rows = []
    for line in raw[1:]:
        if re.match(r"^[\|\-\s:]+$", line.strip()): continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells: rows.append(cells)
    return headers, rows


# ═══════════════════════════════════════════════════════════════════════
# 文档生成器 v2
# ═══════════════════════════════════════════════════════════════════════

class DocxGenerator:
    """专业排版 Word 文档生成器 v2。"""

    def __init__(self, preset: dict[str, Any], meta: dict[str, Any]):
        self.doc = Document()
        self.p = preset
        self.meta = meta
        self._has_cover = False
        self._last_heading1_idx = -1
        self._setup_page(self.doc.sections[0])
        self._setup_styles()

    def _c(self, key: str) -> str:
        """获取颜色 hex。"""
        return self.p["colors"].get(key, "000000")

    def _fc(self, key: str) -> RGBColor:
        return RGBColor.from_string(self._c(key))

    def _fnt(self, role: str) -> str:
        return self.p["fonts"].get(role, self.p["fonts"]["body"])

    def _fsz(self, key: str, default: Pt = Pt(12)) -> Pt:
        return self.p["fonts"].get(key, default)

    # ─── 页面与样式 ──────────────────────────────────────────────────

    def _setup_page(self, section):
        pg = self.p["page"]
        section.page_width = pg["width"]
        section.page_height = pg["height"]
        section.top_margin = pg["margin_top"]
        section.bottom_margin = pg["margin_bottom"]
        section.left_margin = pg["margin_left"]
        section.right_margin = pg["margin_right"]

    def _setup_styles(self):
        """配置 Word 原生样式 Normal + Heading 1/2/3。"""
        body_font = self._fnt("body")
        body_size = self._fsz("body_size", Pt(12))
        heading_font = self._fnt("heading")
        heading_size = self._fsz("heading_size", Pt(16))
        accent = self._fc("accent")
        body_color = self._fc("body")

        # Normal
        ns = self.doc.styles["Normal"]
        ns.font.name = body_font
        ns.font.size = body_size
        ns.font.color.rgb = body_color
        ns.paragraph_format.space_after = Pt(6)
        ns.paragraph_format.line_spacing = 1.5
        self._set_style_eastAsia(ns, body_font)

        # Heading 1
        h1 = self.doc.styles["Heading 1"]
        h1.font.name = heading_font
        h1.font.size = heading_size
        h1.font.bold = True
        h1.font.color.rgb = accent
        self._set_style_eastAsia(h1, heading_font)
        h1.paragraph_format.space_before = Pt(30)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.keep_with_next = True
        # 下方细装饰线
        self._style_bottom_border(h1, self._c("accent"), 4)

        # Heading 2
        h2 = self.doc.styles["Heading 2"]
        h2.font.name = heading_font
        h2.font.size = Pt(int(heading_size.pt * 0.85))
        h2.font.bold = True
        h2.font.color.rgb = RGBColor.from_string("333333")
        self._set_style_eastAsia(h2, heading_font)
        h2.paragraph_format.space_before = Pt(24)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.keep_with_next = True

        # Heading 3
        h3 = self.doc.styles["Heading 3"]
        h3.font.name = heading_font
        h3.font.size = Pt(int(heading_size.pt * 0.72))
        h3.font.bold = True
        h3.font.color.rgb = RGBColor.from_string("555555")
        self._set_style_eastAsia(h3, heading_font)
        h3.paragraph_format.space_before = Pt(18)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.keep_with_next = True

    def _set_style_eastAsia(self, style, font_name: str):
        """设置样式的中文字体。"""
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _style_bottom_border(self, style, color: str, sz: int):
        """给样式添加底部边框。"""
        pPr = style.element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ─── Run 字体设置 ────────────────────────────────────────────────

    def _rf(self, run, font_name: str = None, size: Pt = None, color: RGBColor = None,
            bold: bool = False, italic: bool = False):
        """设置 run 的字体。"""
        fn = font_name or self._fnt("body")
        run.font.name = fn
        run.font.size = size or self._fsz("body_size", Pt(12))
        run.font.color.rgb = color or self._fc("body")
        run.bold = bold
        run.italic = italic
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), fn)

    # ─── 封面页 ──────────────────────────────────────────────────────

    def add_cover(self):
        """生成专业封面页。"""
        if not self.p.get("cover"): return
        self._has_cover = True

        title = self.meta.get("title", "")
        subtitle = self.meta.get("subtitle", "")
        author = self.meta.get("author", "")
        date_str = self.meta.get("date", "") or datetime.now().strftime("%Y年%m月")

        # 封面使用当前 section（之后再加 section break 分页）
        section = self.doc.sections[0]
        # 封面页边距略大，留白更多
        section.top_margin = Cm(4)
        section.bottom_margin = Cm(3)

        bg_color = self._c("cover_bg") or self._c("primary")
        text_color = RGBColor.from_string(self._c("cover_text") or "FFFFFF")

        # ── 顶部色块区域 ──
        # 画一个全宽彩色矩形——用一个大字号空格+shading模拟
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(8)
            # 用段落底纹做色块
            pPr = p._element.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
            pPr.append(shd)

        # ── 标题 ──
        # 留白
        for _ in range(2):
            self.doc.add_paragraph()

        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(16)
        r = p_title.add_run(title)
        self._rf(r, self._fnt("title"), self._fsz("title_size", Pt(26)),
                 self._fc("primary"), bold=True)

        if subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_after = Pt(6)
            r = p_sub.add_run(subtitle)
            self._rf(r, self._fnt("body"), Pt(14), self._fc("body"))

        # ── 分隔装饰线 ──
        p_line = self.doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_before = Pt(20)
        p_line.paragraph_format.space_after = Pt(20)
        pPr = p_line._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{self._c("accent")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # ── 元信息 ──
        meta_texts = []
        if author:
            meta_texts.append(author)
        if date_str:
            meta_texts.append(date_str)
        if self.meta.get("department"):
            meta_texts.append(self.meta["department"])

        for mt in meta_texts:
            pm = self.doc.add_paragraph()
            pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pm.paragraph_format.space_after = Pt(4)
            r = pm.add_run(mt)
            self._rf(r, self._fnt("body"), Pt(11), RGBColor(0x99, 0x99, 0x99))

        # ── 封面分页：插入分节符（下一页） ──
        # 添加分节符，使正文从新页开始，且正文的页眉页脚独立
        new_section = self.doc.add_section()
        self._setup_page(new_section)

        # 正文页的第一页不需要封面信息
        # 如果不需要 TOC 和封面，就留这个 section 放正文
        # 新 section 留在那，后续 build() 会往里加内容
        self._body_section = new_section

    def _body_section_ref(self):
        """返回正文 section（如果有封面则返回封面后的 section）。"""
        if self._has_cover and len(self.doc.sections) > 1:
            return self.doc.sections[1]
        return self.doc.sections[0]

    # ─── TOC ─────────────────────────────────────────────────────────

    def add_toc(self):
        """插入目录域。"""
        if not self.p.get("toc"): return

        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(12)
        r = p_title.add_run("目  录")
        self._rf(r, self._fnt("heading"), Pt(18),
                 self._fc("primary"), bold=True)

        # 装饰线
        p_line = self.doc.add_paragraph()
        pPr = p_line._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="{self._c("accent")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # TOC 域
        p_toc = self.doc.add_paragraph()
        run = p_toc.add_run()
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar_begin)
        run2 = p_toc.add_run()
        instr = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\o "1-3" \\h \\z \\u '
            f'</w:instrText>'
        )
        run2._element.append(instr)
        run3 = p_toc.add_run()
        fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._element.append(fldChar_sep)
        run4 = p_toc.add_run()
        run4._element.append(parse_xml(f'<w:r {nsdecls("w")}><w:t>（在 Word 中右键点击此处 → 更新域 即可生成目录）</w:t></w:r>'))
        run5 = p_toc.add_run()
        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._element.append(fldChar_end)

        # 目录后分页
        self.doc.add_page_break()

    # ─── 标题 ────────────────────────────────────────────────────────

    def add_title(self, text: str):
        """文档主标题（无封面时使用）。"""
        if self._has_cover:
            return

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        self._rf(r, self._fnt("title"), self._fsz("title_size", Pt(22)),
                 self._fc("primary"), bold=True)

        if self.p.get("title_line"):
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(18)
            pPr = p2._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="6" w:color="{self._c("hr_color")}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    def add_subtitle(self, text: str):
        if self._has_cover: return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(text)
        self._rf(r, self._fnt("body"), Pt(14), RGBColor(0x88, 0x88, 0x88))

    # ─── 会议元信息 ──────────────────────────────────────────────────

    def add_meeting_meta(self):
        fields = [
            ("会议时间", self.meta.get("meeting_time") or self.meta.get("date")),
            ("会议地点", self.meta.get("meeting_location")),
            ("主持人", self.meta.get("host")),
            ("记录人", self.meta.get("recorder")),
            ("参会人员", self.meta.get("attendees")),
        ]
        active = [(l, v) for l, v in fields if v]
        if not active: return

        # 元信息放在一个带框的区域
        for label, value in active:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            rl = p.add_run(f"{label}：")
            self._rf(rl, self._fnt("heading"), Pt(12), self._fc("accent"), bold=True)
            rv = p.add_run(value)
            self._rf(rv, self._fnt("body"), Pt(12), self._fc("body"))

        # 分隔线
        p_hr = self.doc.add_paragraph()
        p_hr.paragraph_format.space_before = Pt(8)
        p_hr.paragraph_format.space_after = Pt(12)
        pPr = p_hr._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{self._c("hr_color")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ─── 正文块 ──────────────────────────────────────────────────────

    def add_heading(self, text: str, level: int):
        """使用 Word 原生 Heading 样式。"""
        heading_style = f"Heading {min(level, 3)}"
        p = self.doc.add_paragraph(text, style=heading_style)

        # 如果此 heading 之前有正文内容且是第一个 H1（非封面），前面没有 TOC 则在此位置不够好
        # 保持干净

    def add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        if self.p.get("first_indent"):
            pf.first_line_indent = Cm(0.74)
        pf.space_after = Pt(6)
        self._add_formatted_text(p, text)

    def add_quote(self, text: str):
        """高亮引用块：左侧 accent bar + 浅底色。"""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1.2)
        pf.right_indent = Cm(0.5)
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)

        # 左侧 accent 竖线
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="18" w:space="8" w:color="{self._c("callout_bar")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # 浅底色
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self._c("callout_bg")}" w:val="clear"/>')
        pPr.append(shd)

        r = p.add_run(text)
        self._rf(r, self._fnt("body"), Pt(10.5), RGBColor(0x55, 0x55, 0x55), italic=True)

    def add_bullet_list(self, items: list[str]):
        for item in items:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.5)
            pf.space_after = Pt(3)
            r = p.add_run("•  " + item)
            self._rf(r, self._fnt("body"))

    def add_ordered_list(self, items: list[str]):
        for idx, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.5)
            pf.space_after = Pt(3)
            r = p.add_run(f"{idx}.  {item}")
            self._rf(r, self._fnt("body"))

    def add_table(self, headers: list[str], rows: list[list[str]]):
        """专业表格：细边框、深色表头、交替行色。"""
        if not headers: return

        # 添加表标题
        table_title = self.meta.get("_last_heading_before_table")
        self.doc.add_paragraph()  # 小空行

        cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # 移除默认表格样式，自定义边框
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")} />')

        # 表格边框：细线
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'</w:tblBorders>'
        )
        # 移除旧 borders 如果有
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        tblPr.append(borders)

        hdr_bg = self._c("table_header_bg")
        hdr_text = self._c("table_header_text")
        stripe_bg = self._c("table_stripe")

        # 表头
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # 单元格边距
            tcPr = cell._element.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="40" w:type="dxa"/><w:bottom w:w="40" w:type="dxa"/>'
                f'<w:left w:w="80" w:type="dxa"/><w:right w:w="80" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)
            # 表头底色
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hdr_bg}" w:val="clear"/>')
            tcPr.append(shd)
            r = p.add_run(h)
            self._rf(r, self._fnt("heading"), Pt(11),
                     RGBColor.from_string(hdr_text), bold=True)

        # 数据行
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                if j >= cols: continue
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                tcPr = cell._element.get_or_add_tcPr()
                tcMar = parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    f'<w:top w:w="30" w:type="dxa"/><w:bottom w:w="30" w:type="dxa"/>'
                    f'<w:left w:w="80" w:type="dxa"/><w:right w:w="80" w:type="dxa"/>'
                    f'</w:tcMar>'
                )
                tcPr.append(tcMar)
                if i % 2 == 1:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{stripe_bg}" w:val="clear"/>')
                    tcPr.append(shd)
                r = p.add_run(val)
                self._rf(r, self._fnt("body"), Pt(10.5))

        self.doc.add_paragraph()  # 表后空行

    def add_hr(self):
        """装饰性分隔线。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{self._c("hr_color")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ─── 落款 ────────────────────────────────────────────────────────

    def add_signature(self):
        sign = self.meta.get("signature_unit") or self.meta.get("author") or ""
        date_str = self.meta.get("signature_date") or self.meta.get("date") or ""
        if not sign and not date_str: return

        self.doc.add_paragraph()
        for line in [sign, date_str]:
            if line:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(line)
                self._rf(r, self._fnt("body"))

    # ─── 页眉页脚 ────────────────────────────────────────────────────

    def add_header_footer(self):
        """添加页眉（装饰线 + 标题）和页脚（页码）。"""
        title = self.meta.get("title", "")
        section = self._body_section_ref()

        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(4)
        r = hp.add_run(title if title else "")
        self._rf(r, self._fnt("body"), Pt(8.5), RGBColor(0xAA, 0xAA, 0xAA))

        if self.p.get("header_line"):
            # 页眉下方细线
            pPr = hp._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="2" w:color="{self._c("accent")}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        # 页脚：页码（居中）
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(6)

        # 页脚上方细线
        pPr = fp._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        r = fp.add_run("— ")
        self._rf(r, self._fnt("body"), Pt(9), RGBColor(0xBB, 0xBB, 0xBB))
        # PAGE 域
        for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
            rr = fp.add_run()
            if tag:
                rr._element.append(parse_xml(
                    f'<w:fldChar {nsdecls("w")} w:fldCharType="{tag}"/>'))
            else:
                rr._element.append(parse_xml(
                    f'<w:instrText {nsdecls("w")} xml:space="preserve">{text}</w:instrText>'))
        r2 = fp.add_run(" —")
        self._rf(r2, self._fnt("body"), Pt(9), RGBColor(0xBB, 0xBB, 0xBB))

        # 如果封面页 section 存在，封面不要页眉页脚
        if self._has_cover:
            cs = self.doc.sections[0]
            cs.different_first_page_header_footer = True
            # 封面页的页眉页脚清空
            ch = cs.header
            ch.paragraphs[0].clear()
            cf = cs.footer
            cf.paragraphs[0].clear()

    # ─── 行内格式化 ──────────────────────────────────────────────────

    def _add_formatted_text(self, paragraph, text: str):
        tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
        for tok in tokens:
            if tok.startswith("**") and tok.endswith("**"):
                r = paragraph.add_run(tok[2:-2])
                self._rf(r, bold=True)
            elif tok.startswith("*") and tok.endswith("*") and not tok.startswith("**"):
                r = paragraph.add_run(tok[1:-1])
                self._rf(r, italic=True)
            elif tok.startswith("`") and tok.endswith("`"):
                r = paragraph.add_run(tok[1:-1])
                self._rf(r, font_name="Consolas", size=Pt(10.5),
                         color=RGBColor(0xC7, 0x25, 0x4E))
            elif tok:
                r = paragraph.add_run(tok)
                self._rf(r)

    # ─── 文档属性 ────────────────────────────────────────────────────

    def set_doc_properties(self):
        self.doc.core_properties.title = self.meta.get("title", "")
        self.doc.core_properties.author = self.meta.get("author", "")
        self.doc.core_properties.subject = self.meta.get("subject", "")

    # ─── 主流程 ──────────────────────────────────────────────────────

    def build(self, blocks: list[dict[str, Any]]):
        self.set_doc_properties()

        # 1. 封面
        self.add_cover()

        # 如果无封面，先放标题
        if not self._has_cover:
            title = self.meta.get("title", "")
            if title: self.add_title(title)
            subtitle = self.meta.get("subtitle", "")
            if subtitle: self.add_subtitle(subtitle)

        # 会议元信息
        if self.p["name"] == "会议纪要":
            self.add_meeting_meta()

        # 2. TOC（封面后）
        if self.p.get("toc"):
            self.add_toc()

        # 3. 正文
        last_heading = ""
        for block in blocks:
            btype = block["type"]

            if btype == "heading":
                last_heading = block["text"]
                self.add_heading(block["text"], block["level"])
            elif btype == "paragraph":
                self.add_paragraph(block["text"])
            elif btype == "bullet_list":
                self.add_bullet_list(block["items"])
            elif btype == "ordered_list":
                self.add_ordered_list(block["items"])
            elif btype == "quote":
                self.add_quote(block["text"])
            elif btype == "table":
                headers, rows = parse_table(block["raw"])
                if headers:
                    self.meta["_last_heading_before_table"] = last_heading
                self.add_table(headers, rows)
            elif btype == "hr":
                self.add_hr()

        # 4. 落款
        self.add_signature()

        # 5. 页眉页脚
        self.add_header_footer()

    def save(self, path: str):
        self.doc.save(path)


# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Otto Doc-Writer v2: Markdown + 预设 → 专业排版 .docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
预设: official(公文) | report(报告) | letter(信函) | meeting(会议纪要) | proposal(方案)

示例:
  python create_docx.py input.md output.docx --preset report
  python create_docx.py input.md output.docx -p proposal
        """,
    )
    parser.add_argument("input", help="输入 Markdown 文件")
    parser.add_argument("output", help="输出 .docx 文件")
    parser.add_argument("--preset", "-p", default="report",
                        choices=["official", "report", "letter", "meeting", "proposal"],
                        help="文档预设 (默认: report)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误：找不到 {args.input}"); sys.exit(1)

    text = input_path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)

    if "preset" in meta and args.preset == "report":
        args.preset = meta["preset"]
    if "title" not in meta:
        meta["title"] = input_path.stem

    preset = PRESETS.get(args.preset, PRESETS["report"])
    blocks = parse_markdown_body(body)

    gen = DocxGenerator(preset, meta)
    gen.build(blocks)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    gen.save(str(output_path))

    size_kb = output_path.stat().st_size / 1024
    features = []
    if preset.get("cover"): features.append("封面")
    if preset.get("toc"): features.append("目录")
    print(f"✅ 文档已生成：{output_path}")
    print(f"   预设：{preset['name']}（{' '.join(features) if features else '标准'}）")
    print(f"   大小：{size_kb:.1f} KB · {len(blocks)} 个内容块")


if __name__ == "__main__":
    main()
