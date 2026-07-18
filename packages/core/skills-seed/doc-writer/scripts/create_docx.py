#!/usr/bin/env python3
"""
Otto Doc-Writer v3 — AI 驱动的智能排版引擎

哲学：不做模板。每份文档都应该有自己的视觉语言。
AI 通过 YAML frontmatter 注入设计令牌（颜色/字体/布局），引擎忠实还原。

用法：
  python create_docx.py <input.md> <output.docx>

输入 Markdown 的 YAML frontmatter 中可指定完整设计令牌。不指定时引擎自动推断。
"""
from __future__ import annotations

import json, os, re, sys
from datetime import datetime
from pathlib import Path
from typing import Any

if sys.platform == 'win32':
    try: sys.stdout.reconfigure(encoding='utf-8')
    except Exception: pass

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor, Emu, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误：需要 python-docx。pip install python-docx"); sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════
# 设计令牌解析 — 从 YAML frontmatter 读取 AI 的设计意图
# ═══════════════════════════════════════════════════════════════════════

def parse_design_tokens(meta: dict) -> dict:
    """
    从 frontmatter 解析设计令牌。AI 可写任意键，引擎理解以下：

    色彩：
      primary, accent, body_color, bg_color, muted, table_header_bg,
      table_header_text, table_stripe, callout_bg, callout_bar, hr_color
    字体：
      heading_font, body_font, title_size, h1_size, h2_size, body_size
    排版：
      page_width, page_height, margin_top, margin_bottom, margin_left, margin_right,
      line_spacing, first_indent, title_align, cover, toc
    身份：
      design_name (AI 给这套视觉命名), design_mood
    """
    t: dict[str, Any] = {}

    # ── 颜色 ──
    t["primary"]       = meta.get("primary",       "1B3A5C")
    t["accent"]        = meta.get("accent",        "2E75B6")
    t["body_color"]    = meta.get("body_color",    "333333")
    t["bg_color"]      = meta.get("bg_color",      "")
    t["muted"]         = meta.get("muted",         "888888")
    t["table_header_bg"]   = meta.get("table_header_bg",   t["primary"])
    t["table_header_text"] = meta.get("table_header_text", "FFFFFF")
    t["table_stripe"]      = meta.get("table_stripe",      "F5F7FA")
    t["callout_bg"]        = meta.get("callout_bg",        "F0F4F8")
    t["callout_bar"]       = meta.get("callout_bar",       t["accent"])
    t["hr_color"]          = meta.get("hr_color",          "CCCCCC")
    t["cover_bg"]          = meta.get("cover_bg",          t["primary"])
    t["cover_text"]        = meta.get("cover_text",        "FFFFFF")
    t["header_line_color"] = meta.get("header_line_color", t["accent"])

    # ── 字体 ──
    t["heading_font"] = meta.get("heading_font", "Microsoft YaHei")
    t["body_font"]    = meta.get("body_font",    "SimSun")
    t["title_size"]   = _pt(meta.get("title_size", "26"))
    t["h1_size"]      = _pt(meta.get("h1_size",    "16"))
    t["h2_size"]      = _pt(meta.get("h2_size",    "14"))
    t["body_size"]    = _pt(meta.get("body_size",  "12"))

    # ── 排版 ──
    t["cover"]          = meta.get("cover", "true") != "false"
    t["toc"]            = meta.get("toc",   "true")  == "true"
    t["first_indent"]   = meta.get("first_indent",  "true") != "false"
    t["title_align"]    = meta.get("title_align",    "center")
    t["line_spacing"]   = float(meta.get("line_spacing", "1.5"))
    t["page_width"]     = Cm(float(meta.get("page_width", "21")))
    t["page_height"]    = Cm(float(meta.get("page_height", "29.7")))
    t["margin_top"]     = Cm(float(meta.get("margin_top", "2.54")))
    t["margin_bottom"]  = Cm(float(meta.get("margin_bottom", "2.54")))
    t["margin_left"]    = Cm(float(meta.get("margin_left", "3.18")))
    t["margin_right"]   = Cm(float(meta.get("margin_right", "3.18")))

    # ── 身份 ──
    t["design_name"] = meta.get("design_name", "")
    t["design_mood"] = meta.get("design_mood", "")

    return t


def _pt(val) -> Pt:
    return Pt(float(str(val).replace("pt", "")))


# ═══════════════════════════════════════════════════════════════════════
# Markdown 解析
# ═══════════════════════════════════════════════════════════════════════

def parse_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    meta: dict[str, Any] = {}
    body = text.strip()
    if body.startswith("---"):
        parts = body.split("---", 2)
        if len(parts) >= 3:
            for line in parts[1].strip().split("\n"):
                line = line.strip()
                if ":" in line and not line.startswith("#"):
                    key, _, val = line.partition(":")
                    meta[key.strip()] = val.strip().strip('"').strip("'")
            body = parts[2].strip()
    return meta, body


def parse_markdown_body(text: str) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    lines = text.split("\n"); i = 0
    table_buf: list[str] = []; in_table = False

    def flush():
        nonlocal table_buf, in_table
        if table_buf:
            blocks.append({"type": "table", "raw": table_buf})
            table_buf = []; in_table = False

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buf.append(line); in_table = True; i += 1; continue
        elif in_table:
            flush(); continue
        if not line.strip(): i += 1; continue

        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            flush()
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1; continue

        m = re.match(r"^[-*+]\s+(.+)$", line)
        if m:
            flush()
            items = []
            while i < len(lines) and re.match(r"^[-*+]\s+(.+)$", lines[i]):
                items.append(re.match(r"^[-*+]\s+(.+)$", lines[i]).group(1).strip())
                i += 1
            blocks.append({"type": "bullet_list", "items": items}); continue

        m = re.match(r"^\d+[.)]\s+(.+)$", line)
        if m:
            flush()
            items = []
            while i < len(lines) and re.match(r"^\d+[.)]\s+(.+)$", lines[i]):
                items.append(re.match(r"^\d+[.)]\s+(.+)$", lines[i]).group(1).strip())
                i += 1
            blocks.append({"type": "ordered_list", "items": items}); continue

        if line.startswith("> "):
            flush()
            qlines = []
            while i < len(lines) and lines[i].startswith("> "):
                qlines.append(lines[i][2:].strip())
                i += 1
            blocks.append({"type": "quote", "text": "\n".join(qlines)}); continue

        if line.strip() in ("---", "***", "___"):
            flush()
            blocks.append({"type": "hr"}); i += 1; continue

        flush()
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and \
              not re.match(r"^[-*+]\s+", lines[i]) and not re.match(r"^\d+[.)]\s+", lines[i]) and \
              not lines[i].startswith("|") and not lines[i].startswith("> ") and \
              lines[i].strip() not in ("---","***","___"):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"type": "paragraph", "text": "\n".join(para_lines)})

    flush()
    return blocks


def parse_table(raw: list[str]) -> tuple[list[str], list[list[str]]]:
    if len(raw) < 2: return [], []
    headers = [c.strip() for c in raw[0].strip("|").split("|")]
    rows = []
    for line in raw[1:]:
        if re.match(r"^[\|\-\s:]+$", line.strip()): continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells: rows.append(cells)
    return headers, rows


# ═══════════════════════════════════════════════════════════════════════
# 文档渲染引擎 — 纯渲染，不选模板
# ═══════════════════════════════════════════════════════════════════════

class DocxRenderer:
    """设计令牌驱动的文档渲染器。"""

    def __init__(self, tokens: dict[str, Any], meta: dict[str, Any]):
        self.doc = Document()
        self.t = tokens   # 设计令牌
        self.m = meta     # 文档元数据
        self._has_cover = False
        self._setup_page(self.doc.sections[0])
        self._setup_styles()

    # ── 令牌读取 ──

    def _tc(self, key: str) -> str:  return self.t.get(key, "000000")
    def _clr(self, key: str) -> RGBColor: return RGBColor.from_string(self._tc(key))
    def _fnt(self, key: str) -> str: return self.t.get(key + "_font" if key in ("heading","body") else key, "SimSun")
    def _fsz(self, key: str) -> Pt:  return self.t.get(key + "_size", Pt(12)) if key in ("h1","h2","body","title") else Pt(12)

    # ── 页面 ──

    def _setup_page(self, section):
        section.page_width  = self.t["page_width"]
        section.page_height = self.t["page_height"]
        section.top_margin    = self.t["margin_top"]
        section.bottom_margin = self.t["margin_bottom"]
        section.left_margin   = self.t["margin_left"]
        section.right_margin  = self.t["margin_right"]

    def _setup_styles(self):
        bf = self.t["body_font"]
        bs = self.t["body_size"]
        hf = self.t["heading_font"]

        ns = self.doc.styles["Normal"]
        ns.font.name = bf; ns.font.size = bs
        ns.font.color.rgb = self._clr("body_color")
        ns.paragraph_format.space_after = Pt(6)
        ns.paragraph_format.line_spacing = self.t["line_spacing"]
        self._eastAsia(ns, bf)

        for lvl, (sz, color_key) in enumerate([
            (self.t["h1_size"], "primary"),
            (self.t["h2_size"], "body_color"),
            (Pt(self.t["h2_size"].pt * 0.85), "muted"),
        ], 1):
            s = self.doc.styles[f"Heading {lvl}"]
            s.font.name = hf; s.font.size = sz; s.font.bold = True
            s.font.color.rgb = self._clr(color_key)
            self._eastAsia(s, hf)
            s.paragraph_format.space_before = Pt(30 if lvl == 1 else 24 if lvl == 2 else 18)
            s.paragraph_format.space_after  = Pt(12 if lvl == 1 else 8  if lvl == 2 else 6)
            s.paragraph_format.keep_with_next = True
            if lvl == 1:
                self._style_bottom_border(s, self._tc("accent"), 6)

    def _eastAsia(self, style, name):
        rPr = style.element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rf)
        rf.set(qn("w:eastAsia"), name)

    def _style_bottom_border(self, style, color, sz):
        pPr = style.element.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/>'
            f'</w:pBdr>'))

    # ── Run 排版 ──

    def _run(self, run, *, name=None, size=None, color=None, bold=False, italic=False):
        fn = name or self.t["body_font"]
        run.font.name = fn
        run.font.size = size or self.t["body_size"]
        run.font.color.rgb = color or self._clr("body_color")
        run.bold = bold; run.italic = italic
        rPr = run._element.get_or_add_rPr()
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rf)
        rf.set(qn("w:eastAsia"), fn)

    # ── 封面 ──

    def add_cover(self):
        if not self.t.get("cover"): return
        self._has_cover = True
        section = self.doc.sections[0]
        section.top_margin = Cm(4); section.bottom_margin = Cm(3)

        title = self.m.get("title", "")
        subtitle = self.m.get("subtitle", "")
        author = self.m.get("author", "")
        date_str = self.m.get("date", "") or datetime.now().strftime("%Y年%m月")
        primary = self._tc("cover_bg") or self._tc("primary")
        accent  = self._tc("accent")
        muted   = self._tc("muted")

        # 顶部色块
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(8)
            pPr = p._element.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{primary}" w:val="clear"/>'))

        for _ in range(3): self.doc.add_paragraph()

        # 标题
        pt = self.doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_after = Pt(16)
        r = pt.add_run(title)
        self._run(r, name=self.t["heading_font"], size=self.t["title_size"],
                  color=self._clr("primary"), bold=True)

        if subtitle:
            ps = self.doc.add_paragraph()
            ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ps.paragraph_format.space_after = Pt(6)
            r = ps.add_run(subtitle)
            self._run(r, size=Pt(14), color=self._clr("body_color"))

        # 装饰线
        pl = self.doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_before = Pt(24)
        pl.paragraph_format.space_after = Pt(24)
        pl._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{accent}"/>'
            f'</w:pBdr>'))

        # 元信息
        meta_items = []
        if author: meta_items.append(author)
        if date_str: meta_items.append(date_str)
        if self.m.get("department"): meta_items.append(self.m["department"])
        for mt in meta_items:
            pm = self.doc.add_paragraph()
            pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pm.paragraph_format.space_after = Pt(4)
            self._run(pm.add_run(mt), size=Pt(11), color=RGBColor.from_string("999999"))

        # 分节
        new_sec = self.doc.add_section()
        self._setup_page(new_sec)

    def _body_section(self):
        if self._has_cover and len(self.doc.sections) > 1:
            return self.doc.sections[1]
        return self.doc.sections[0]

    # ── TOC ──

    def add_toc(self):
        if not self.t.get("toc"): return

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        self._run(p.add_run("目  录"), name=self.t["heading_font"],
                  size=Pt(18), color=self._clr("primary"), bold=True)

        pl = self.doc.add_paragraph()
        pl._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="{self._tc("accent")}"/>'
            f'</w:pBdr>'))

        ptoc = self.doc.add_paragraph()
        r = ptoc.add_run()
        r._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        r2 = ptoc.add_run()
        r2._element.append(parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
        r3 = ptoc.add_run()
        r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        r4 = ptoc.add_run("（在 Word 中右键 → 更新域 即可生成目录）")
        self._run(r4, size=Pt(10), color=RGBColor.from_string("AAAAAA"), italic=True)
        r5 = ptoc.add_run()
        r5._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

        self.doc.add_page_break()

    # ── 标题（无封面时） ──

    def add_title(self, text: str):
        if self._has_cover: return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if self.t["title_align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        self._run(p.add_run(text), name=self.t["heading_font"],
                  size=self.t["title_size"], color=self._clr("primary"), bold=True)

        # 装饰线
        p2 = self.doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        p2._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="6" w:color="{self._tc("accent")}"/>'
            f'</w:pBdr>'))

    def add_subtitle(self, text: str):
        if self._has_cover: return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if self.t["title_align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(24)
        self._run(p.add_run(text), size=Pt(14), color=RGBColor.from_string("999999"))

    # ── 正文块 ──

    def add_heading(self, text: str, level: int):
        self.doc.add_paragraph(text, style=f"Heading {min(level, 3)}")

    def add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        if self.t.get("first_indent"): p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(6)
        self._fmt_text(p, text)

    def add_quote(self, text: str):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(1.2); pf.right_indent = Cm(0.5)
        pf.space_before = Pt(14); pf.space_after = Pt(14)
        pPr = p._element.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="18" w:space="8" w:color="{self._tc("callout_bar")}"/>'
            f'</w:pBdr>'))
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self._tc("callout_bg")}" w:val="clear"/>'))
        self._run(p.add_run(text), size=Pt(10.5), color=RGBColor.from_string("666666"), italic=True)

    def add_bullet_list(self, items: list[str]):
        for item in items:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-0.5)
            pf.space_after = Pt(3)
            self._run(p.add_run("•  " + item))

    def add_ordered_list(self, items: list[str]):
        for idx, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0); pf.first_line_indent = Cm(-0.5)
            pf.space_after = Pt(3)
            self._run(p.add_run(f"{idx}.  {item}"))

    def add_table(self, headers: list[str], rows: list[list[str]]):
        if not headers: return
        self.doc.add_paragraph()
        cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")} />')
        for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
            f'</w:tblBorders>'))

        hdr_bg   = self._tc("table_header_bg")
        hdr_text = self._tc("table_header_text")
        stripe   = self._tc("table_stripe")

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]; cell.text = ""
            cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(5)
            cp.paragraph_format.space_after = Pt(5)
            cell._element.get_or_add_tcPr().append(parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="50" w:type="dxa"/><w:bottom w:w="50" w:type="dxa"/>'
                f'<w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/>'
                f'</w:tcMar>'))
            cell._element.get_or_add_tcPr().append(parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{hdr_bg}" w:val="clear"/>'))
            self._run(cp.add_run(h), name=self.t["heading_font"], size=Pt(10.5),
                      color=RGBColor.from_string(hdr_text), bold=True)

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                if j >= cols: continue
                cell = table.rows[i + 1].cells[j]; cell.text = ""
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                cell._element.get_or_add_tcPr().append(parse_xml(
                    f'<w:tcMar {nsdecls("w")}>'
                    f'<w:top w:w="40" w:type="dxa"/><w:bottom w:w="40" w:type="dxa"/>'
                    f'<w:left w:w="100" w:type="dxa"/><w:right w:w="100" w:type="dxa"/>'
                    f'</w:tcMar>'))
                if i % 2 == 1:
                    cell._element.get_or_add_tcPr().append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{stripe}" w:val="clear"/>'))
                self._run(cp.add_run(val), size=Pt(10.5))
        self.doc.add_paragraph()

    def add_hr(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        p._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{self._tc("hr_color")}"/>'
            f'</w:pBdr>'))

    # ── 签名 ──

    def add_signature(self):
        sign = self.m.get("signature_unit") or self.m.get("author") or ""
        date_str = self.m.get("signature_date") or self.m.get("date") or ""
        if not sign and not date_str: return
        self.doc.add_paragraph()
        for line in [sign, date_str]:
            if line:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(4)
                self._run(p.add_run(line))

    # ── 页眉页脚 ──

    def add_header_footer(self):
        title = self.m.get("title", "")
        section = self._body_section()

        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.paragraph_format.space_after = Pt(4)
        self._run(hp.add_run(title or ""), size=Pt(8.5), color=RGBColor.from_string("BBBBBB"))

        hcol = self._tc("header_line_color")
        if hcol and hcol != "none":
            hp._element.get_or_add_pPr().append(parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="2" w:color="{hcol}"/>'
                f'</w:pBdr>'))

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(8)
        fp._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="4" w:color="E0E0E0"/>'
            f'</w:pBdr>'))

        self._run(fp.add_run("— "), size=Pt(9), color=RGBColor.from_string("CCCCCC"))
        for tag, text in [("begin", None), (None, " PAGE "), ("end", None)]:
            rr = fp.add_run()
            rr._element.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="{tag}"/>' if tag else
                f'<w:instrText {nsdecls("w")} xml:space="preserve">{text}</w:instrText>'))
        self._run(fp.add_run(" —"), size=Pt(9), color=RGBColor.from_string("CCCCCC"))

        if self._has_cover:
            cs = self.doc.sections[0]
            cs.different_first_page_header_footer = True
            cs.header.paragraphs[0].clear()
            cs.footer.paragraphs[0].clear()

    # ── 行内格式化 ──

    def _fmt_text(self, paragraph, text: str):
        tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
        for tok in tokens:
            if tok.startswith("**") and tok.endswith("**"):
                self._run(paragraph.add_run(tok[2:-2]), bold=True)
            elif tok.startswith("*") and tok.endswith("*") and not tok.startswith("**"):
                self._run(paragraph.add_run(tok[1:-1]), italic=True)
            elif tok.startswith("`") and tok.endswith("`"):
                self._run(paragraph.add_run(tok[1:-1]), name="Consolas", size=Pt(10),
                          color=RGBColor(0xC7, 0x25, 0x4E))
            elif tok:
                self._run(paragraph.add_run(tok))

    # ── 文档属性 ──

    def _set_props(self):
        self.doc.core_properties.title = self.m.get("title", "")
        self.doc.core_properties.author = self.m.get("author", "")
        self.doc.core_properties.subject = self.m.get("subject", "")

    # ── 主流程 ──

    def build(self, blocks: list[dict[str, Any]]):
        self._set_props()
        self.add_cover()

        if not self._has_cover:
            title = self.m.get("title", "")
            if title: self.add_title(title)
            sub = self.m.get("subtitle", "")
            if sub: self.add_subtitle(sub)

        if self.t.get("toc"): self.add_toc()

        for block in blocks:
            t = block["type"]
            if   t == "heading":      self.add_heading(block["text"], block["level"])
            elif t == "paragraph":    self.add_paragraph(block["text"])
            elif t == "bullet_list":  self.add_bullet_list(block["items"])
            elif t == "ordered_list": self.add_ordered_list(block["items"])
            elif t == "quote":        self.add_quote(block["text"])
            elif t == "table":
                h, r = parse_table(block["raw"])
                self.add_table(h, r)
            elif t == "hr":           self.add_hr()

        self.add_signature()
        self.add_header_footer()

    def save(self, path: str):
        self.doc.save(path)


# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Otto Doc-Writer v3: AI 驱动排版引擎",
        epilog="AI 通过 YAML frontmatter 注入设计令牌。不指定时引擎自动推断。")
    parser.add_argument("input", help="输入 Markdown 文件")
    parser.add_argument("output", help="输出 .docx 文件")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"错误：找不到 {args.input}"); sys.exit(1)

    text = input_path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)
    tokens = parse_design_tokens(meta)

    if "title" not in meta:
        meta["title"] = input_path.stem

    blocks = parse_markdown_body(body)
    gen = DocxRenderer(tokens, meta)
    gen.build(blocks)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    gen.save(str(output_path))

    size_kb = output_path.stat().st_size / 1024
    dn = tokens.get("design_name", "")
    dm = tokens.get("design_mood", "")
    tag = ""
    if dn:  tag += dn
    if dm:  tag += (" · " + dm) if tag else dm
    print(f"✅ 文档已生成：{output_path}")
    print(f"   视觉语言：{tag}" if tag else f"   视觉语言：自动推断")
    print(f"   大小：{size_kb:.1f} KB · {len(blocks)} 个内容块")


if __name__ == "__main__":
    main()
